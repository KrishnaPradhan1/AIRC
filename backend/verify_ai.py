import requests
import os

# Login to get token
login_url = 'http://localhost:5000/api/auth/login'
upload_url = 'http://localhost:5000/api/resume/upload'

email = 'stu2@test.com'
password = 'password123'

session = requests.Session()

print(f"Logging in as {email}...")
try:
    response = session.post(login_url, json={'email': email, 'password': password})
    if response.status_code == 200:
        token = response.json()['access_token']
        headers = {'Authorization': f'Bearer {token}'}
        print("Login successful.")
    else:
        print(f"Login failed: {response.text}")
        exit(1)
except Exception as e:
    print(f"Connection failed: {e}")
    exit(1)

# Upload Resume
file_path = 'dummy_resume.docx'
if not os.path.exists(file_path):
    # Create a dummy docx if not exists
    import docx
    doc = docx.Document()
    doc.add_heading('Sameer Kumar Sahu', 0)
    doc.add_paragraph('Python Developer | Data Scientist')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, SQL, React, Flask, Machine Learning')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Engineer at Tech Corp (2020-Present). Built AI models.')
    doc.save(file_path)
    print("Created dummy_resume.docx")

print("Uploading resume and waiting for AI analysis...")
with open(file_path, 'rb') as f:
    files = {'resume': f}
    try:
        response = session.post(upload_url, headers=headers, files=files)
        if response.status_code == 201:
            print("Upload successful!")
            data = response.json()
            print("Response Data:", data)
            if 'analysis' in data:
                print("\n--- AI Analysis Result ---")
                print(data['analysis'])
                print("--------------------------")
            else:
                print("Warning: No analysis data in response.")
        else:
            print(f"Upload failed: {response.status_code} - {response.text}")
    except Exception as e:
        print(f"Upload request failed: {e}")
